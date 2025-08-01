#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
مولد مستند ورد RTL فارسی
برای تبدیل فایل‌های مخزن translation به فایل ورد با قالب‌بندی RTL صحیح
"""

import os
import re
import base64
import requests
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from PIL import Image
import io
from time import sleep

# اطلاعات مخزن GitHub
GITHUB_OWNER = "ehsan4232"
GITHUB_REPO = "translation"
GITHUB_API_BASE = "https://api.github.com/repos"

def setup_rtl_paragraph(paragraph):
    """تنظیم فرمت RTL برای پاراگراف"""
    try:
        pPr = paragraph._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    except Exception:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return paragraph

def setup_rtl_run(run):
    """تنظیم فرمت RTL برای قسمت متن"""
    try:
        rPr = run._element.get_or_add_rPr()
        rtl = OxmlElement('w:rtl')
        rtl.set(qn('w:val'), '1')
        rPr.append(rtl)
        run.font.name = 'B Nazanin'
        run.font.size = Pt(12)
    except Exception:
        try:
            run.font.name = 'Tahoma'
            run.font.size = Pt(12)
        except Exception:
            pass
    return run

def fix_farsi_text(text):
    """اصلاح متن فارسی برای نمایش صحیح RTL"""
    if not text or not text.strip():
        return text
    
    try:
        # اضافه کردن نشانه RTL در ابتدای متن فارسی
        if re.match(r'^[\u0600-\u06FF]', text.strip()):
            text = '\u200F' + text  # Right-to-Left Mark
        
        # مدیریت bullet point ها و شماره‌گذاری
        text = re.sub(r'^[\s]*[-•·]\s*', '• ', text, flags=re.MULTILINE)
        text = re.sub(r'^[\s]*(\d+)[\.\)]\s*', r'\1. ', text, flags=re.MULTILINE)
        
        # اضافه کردن word joiner برای کلمات انگلیسی درون متن فارسی
        text = re.sub(r'([\u0600-\u06FF\s])([a-zA-Z0-9]+)([\u0600-\u06FF\s])', 
                      r'\1\u2060\2\u2060\3', text)
        
        # مدیریت علائم نگارشی در محیط RTL
        text = text.replace('(', '\u200E(\u200F').replace(')', '\u200E)\u200F')
        text = text.replace('[', '\u200E[\u200F').replace(']', '\u200E]\u200F')
        text = text.replace('{', '\u200E{\u200F').replace('}', '\u200E}\u200F')
        
    except Exception as e:
        print(f"    هشدار: خطا در قالب‌بندی RTL: {e}")
    
    return text

def natural_sort_key(filename):
    """ایجاد کلید مرتب‌سازی طبیعی"""
    parts = re.split(r'(\d+)', filename)
    return [int(part) if part.isdigit() else part.lower() for part in parts]

def get_repo_files():
    """دریافت تمام فایل‌ها از مخزن GitHub"""
    url = f"{GITHUB_API_BASE}/{GITHUB_OWNER}/{GITHUB_REPO}/contents"
    
    try:
        response = requests.get(url)
        response.raise_for_status()
        files = response.json()
        
        # فیلتر کردن فایل‌های .txt و .png
        target_files = []
        for file in files:
            if file['type'] == 'file' and file['name'].lower().endswith(('.txt', '.png')):
                target_files.append(file)
        
        # مرتب‌سازی طبیعی فایل‌ها
        target_files.sort(key=lambda x: natural_sort_key(x['name']))
        
        return target_files
        
    except Exception as e:
        print(f"خطا در دریافت فایل‌های مخزن: {e}")
        return []

def download_file_content(file_info):
    """دانلود محتوای فایل از GitHub"""
    try:
        if file_info['name'].endswith('.txt'):
            # برای فایل‌های متنی، استفاده از محتوای API (base64 encoded)
            if 'content' in file_info:
                content = base64.b64decode(file_info['content']).decode('utf-8')
                return content
            else:
                # راه جایگزین: دانلود از URL خام
                response = requests.get(file_info['download_url'])
                response.raise_for_status()
                return response.text
        
        elif file_info['name'].endswith('.png'):
            # برای تصاویر، دانلود محتوای باینری
            response = requests.get(file_info['download_url'])
            response.raise_for_status()
            return response.content
            
    except Exception as e:
        print(f"    خطا در دانلود {file_info['name']}: {e}")
        return None

def process_github_repo_to_word(output_folder="output_documents"):
    """دانلود فایل‌ها از مخزن GitHub و ایجاد مستند ورد"""
    output_path = Path(output_folder)
    output_path.mkdir(exist_ok=True)
    
    print(f"🔍 در حال دریافت فایل‌ها از مخزن GitHub: {GITHUB_OWNER}/{GITHUB_REPO}")
    
    # دریافت تمام فایل‌ها از مخزن
    files = get_repo_files()
    
    if not files:
        raise ValueError("هیچ فایلی در مخزن یافت نشد!")
    
    print(f"📁 {len(files)} فایل برای پردازش یافت شد")
    
    # ایجاد مستند ورد
    doc = Document()
    
    # تنظیم جهت مستند به RTL
    try:
        section = doc.sections[0]
        sectPr = section._sectPr
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        sectPr.append(bidi)
    except Exception:
        pass
    
    # اضافه کردن عنوان
    title = doc.add_heading('مستندات طراحی سیستم', level=1)
    setup_rtl_paragraph(title)
    
    # اضافه کردن زیرعنوان با اطلاعات مخزن
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(f'منبع: مخزن GitHub - {GITHUB_OWNER}/{GITHUB_REPO}')
    setup_rtl_paragraph(subtitle)
    setup_rtl_run(subtitle_run)
    
    doc.add_paragraph()  # خط خالی
    
    # پردازش هر فایل
    for i, file_info in enumerate(files, 1):
        print(f"📄 در حال پردازش {i}/{len(files)}: {file_info['name']}")
        
        # اضافه کردن جداکننده فایل
        separator = doc.add_paragraph()
        separator_run = separator.add_run('=' * 60)
        separator_run.font.name = 'Courier New'
        setup_rtl_paragraph(separator)
        
        # اضافه کردن نام فایل به عنوان سرتیتر
        file_heading = doc.add_heading(f"فایل {i}: {file_info['name']}", level=2)
        setup_rtl_paragraph(file_heading)
        
        # دانلود و پردازش محتوای فایل
        content = download_file_content(file_info)
        
        if content is None:
            # خطا در دانلود فایل
            error_para = doc.add_paragraph()
            error_run = error_para.add_run("❌ خطا در دانلود فایل")
            setup_rtl_paragraph(error_para)
            setup_rtl_run(error_run)
            continue
        
        if file_info['name'].endswith('.txt'):
            # پردازش فایل متنی
            try:
                print(f"  ✍️  پردازش متن فارسی...")
                
                if content.strip():
                    # پاک‌سازی و اصلاح متن برای نمایش RTL
                    lines = content.split('\n')
                    
                    for line in lines:
                        line = line.strip()
                        if line:
                            try:
                                # اعمال قالب‌بندی RTL
                                fixed_line = fix_farsi_text(line)
                                
                                # ایجاد پاراگراف
                                para = doc.add_paragraph()
                                run = para.add_run(fixed_line)
                                
                                # اعمال قالب‌بندی RTL
                                setup_rtl_paragraph(para)
                                setup_rtl_run(run)
                                
                            except Exception as line_error:
                                print(f"    ⚠️  هشدار: مشکل در قالب‌بندی خط: {line_error}")
                                # راه جایگزین: اضافه کردن خط بدون قالب‌بندی پیشرفته
                                para = doc.add_paragraph()
                                run = para.add_run(line)
                                setup_rtl_paragraph(para)
                                setup_rtl_run(run)
                        else:
                            # خط خالی
                            doc.add_paragraph()
                    
                    print(f"  ✅ فایل متنی با موفقیت پردازش شد")
                
                else:
                    # فایل خالی
                    empty_para = doc.add_paragraph()
                    empty_run = empty_para.add_run("(فایل خالی)")
                    setup_rtl_paragraph(empty_para)
                    setup_rtl_run(empty_run)
                    
            except Exception as e:
                print(f"  ❌ خطا در پردازش متن: {e}")
                error_para = doc.add_paragraph()
                error_run = error_para.add_run(f"خطا در پردازش متن: نمی‌توان محتوا را نمایش داد")
                setup_rtl_paragraph(error_para)
                setup_rtl_run(error_run)
        
        elif file_info['name'].endswith('.png'):
            # پردازش فایل تصویر
            try:
                print(f"  🖼️  پردازش تصویر...")
                
                # ایجاد یک BytesIO object از محتوای دانلود شده
                image_stream = io.BytesIO(content)
                
                # بررسی ابعاد تصویر
                with Image.open(image_stream) as img:
                    width, height = img.size
                    max_width = 6  # اینچ
                    
                    if width > height:
                        # جهت landscape
                        img_width = min(max_width, width / 100)
                        img_height = img_width * height / width
                    else:
                        # جهت portrait  
                        img_height = min(max_width, height / 100)
                        img_width = img_height * width / height
                
                # بازنشانی موقعیت stream
                image_stream.seek(0)
                
                # اضافه کردن تصویر به مستند
                para = doc.add_paragraph()
                setup_rtl_paragraph(para)
                run = para.add_run()
                run.add_picture(image_stream, width=Inches(img_width))
                
                print(f"  ✅ تصویر با موفقیت اضافه شد")
                
            except Exception as e:
                print(f"  ❌ خطا در پردازش تصویر: {e}")
                error_para = doc.add_paragraph()
                error_run = error_para.add_run(f"خطا در بارگذاری تصویر: {str(e)}")
                setup_rtl_paragraph(error_para)
                setup_rtl_run(error_run)
        
        # اضافه کردن فاصله بعد از هر فایل
        doc.add_paragraph()
        
        # کمی توقف برای جلوگیری از محدودیت نرخ
        sleep(0.1)
    
    # ذخیره مستند
    output_file = output_path / f"{GITHUB_REPO}_مستندات_طراحی_سیستم_RTL.docx"
    doc.save(str(output_file))
    
    print(f"\n🎉 مستند ورد با موفقیت ایجاد شد!")
    print(f"📁 ذخیره شده در: {output_file}")
    print(f"📊 {len(files)} فایل از مخزن GitHub پردازش شد")
    
    return output_file

def main():
    """تابع اصلی"""
    print("🚀 مولد مستند ورد RTL فارسی")
    print("=" * 60)
    print(f"📂 منبع: https://github.com/{GITHUB_OWNER}/{GITHUB_REPO}")
    print()
    
    try:
        output_file = process_github_repo_to_word()
        
        print(f"\n✨ ویژگی‌های مستند:")
        print("   ✓ قالب‌بندی RTL صحیح برای متن فارسی")
        print("   ✓ مدیریت صحیح کلمات انگلیسی درون جملات فارسی")
        print("   ✓ bullet point ها و شماره‌گذاری صحیح")
        print("   ✓ تمام تصاویر با اندازه مناسب")
        print("   ✓ مرتب‌سازی طبیعی فایل‌ها (1، 2، 3، ...)")
        print("   ✓ دانلود مستقیم از مخزن GitHub")
        print(f"\n📖 فایل مستند را باز کنید: {output_file}")
        
    except Exception as e:
        print(f"❌ خطا: {e}")
        print("لطفاً اتصال اینترنت خود را بررسی کرده و دوباره تلاش کنید.")

if __name__ == "__main__":
    main()
